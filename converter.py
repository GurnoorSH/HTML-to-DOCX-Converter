from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_html_to_docx(html_content):
    doc = Document()
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(True):
        if element.name == 'h1':
            p = doc.add_heading(level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(element.text)
            run.bold = True
            run.font.size = Pt(24)
        elif element.name == 'h2':
            p = doc.add_heading(level=2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(element.text)
            run.bold = True
            run.font.size = Pt(18)
        elif element.name == 'p':
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    p.add_run(child.text).bold = True
                elif child.name == 'em' or child.name == 'i':
                    p.add_run(child.text).italic = True
                elif child.name == 'u':
                    p.add_run(child.text).underline = True
                elif child.name == 'a':
                    run = p.add_run(child.text)
                    run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                    run.underline = True
                else:
                    p.add_run(str(child))
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == 'table':
            table = doc.add_table(rows=0, cols=len(element.find_all('th')))
            table.style = 'Table Grid'
            for row in element.find_all('tr'):
                cells = row.find_all(['th', 'td'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell.text
    return doc
