from flask import Flask, request, send_file
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

app = Flask(__name__, static_folder='frontend/build', static_url_path='/')

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = (0, 0, 255)
    r.font.underline = True

    return hyperlink

def handle_tag(tag, doc, parent_element):
    if tag.name == 'p':
        p = doc.add_paragraph()
        for child in tag.children:
            if child.name == 'a':
                add_hyperlink(p, child.get_text(strip=True), child.get('href'))
            else:
                p.add_run(child.get_text(strip=True))
    elif tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag.name[1])
        p = doc.add_heading(level=level)
        p.add_run(tag.get_text(strip=True))
    elif tag.name == 'ul':
        for li in tag.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            handle_tag(li, doc, p)
    elif tag.name == 'ol':
        for li in tag.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Number')
            handle_tag(li, doc, p)
    elif tag.name == 'li':
        for child in tag.children:
            if child.name:
                handle_tag(child, doc, parent_element)
            else:
                parent_element.add_run(child.string)
    elif tag.name == 'table':
        table = doc.add_table(rows=0, cols=len(tag.find('tr').find_all(['th', 'td'])))
        table.style = 'Table Grid'
        for row in tag.find_all('tr'):
            row_cells = table.add_row().cells
            for i, cell in enumerate(row.find_all(['th', 'td'])):
                row_cells[i].text = cell.get_text(strip=True)
    elif tag.name == 'a':
         add_hyperlink(parent_element, tag.get_text(strip=True), tag.get('href'))
    else:
        if tag.string:
            parent_element.add_run(tag.string)


def convert_html_to_docx(html_content, docx_filename):
    doc = Document()
    soup = BeautifulSoup(html_content, 'lxml')

    for element in soup.body.children:
        if element.name:
            handle_tag(element, doc, doc)

    doc.save(docx_filename)

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'htmlFile' not in request.files:
        return 'No file part', 400

    file = request.files['htmlFile']

    if file.filename == '':
        return 'No selected file', 400

    if file:
        html_content = file.read().decode('utf-8')
        output_filename = 'converted.docx'
        convert_html_to_docx(html_content, output_filename)
        return send_file(output_filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
