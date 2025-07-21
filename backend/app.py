from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import uuid
import time
import threading
from datetime import datetime

app = Flask(__name__)
CORS(app)  # Enable CORS for frontend communication

# In-memory job storage (in production, use a database)
jobs = {}
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)
    r.font.color.rgb = RGBColor(0, 0, 255)
    r.font.underline = True

    return hyperlink

def apply_styles(run, styles):
    if 'font-size' in styles:
        run.font.size = Pt(float(styles['font-size'].replace('pt', '')))
    if 'font-weight' in styles and styles['font-weight'] == 'bold':
        run.font.bold = True
    if 'font-style' in styles and styles['font-style'] == 'italic':
        run.font.italic = True
    if 'color' in styles:
        color = styles['color'].replace('#', '')
        run.font.color.rgb = RGBColor.from_string(color)

def handle_tag(tag, doc, parent_element=None):
    styles = {}
    if tag.get('style'):
        styles = {
            prop.strip(): val.strip()
            for prop, val in (
                style.split(':') for style in tag.get('style').split(';') if ':' in style
            )
        }

    if tag.name == 'p':
        p = doc.add_paragraph()
        if 'text-align' in styles:
            align = styles['text-align']
            if align == 'center':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            elif align == 'right':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            elif align == 'justify':
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for child in tag.children:
            if child.name == 'a':
                add_hyperlink(p, child.get_text(strip=True), child.get('href'))
            elif child.name:
                handle_tag(child, doc, p)
            elif child.string and child.string.strip():
                run = p.add_run(child.string.strip())
                apply_styles(run, styles)

    elif tag.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        level = int(tag.name[1])
        p = doc.add_heading(level=level)
        text_content = tag.get_text(strip=True)
        if text_content:
            run = p.add_run(text_content)
            apply_styles(run, styles)

    elif tag.name == 'ul':
        for li in tag.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            handle_tag(li, doc, p)
    elif tag.name == 'ol':
        for li in tag.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Number')
            handle_tag(li, doc, p)
    elif tag.name == 'li':
        for child in tag.children:
            if child.name:
                handle_tag(child, doc, parent_element)
            elif child.string and child.string.strip() and parent_element:
                run = parent_element.add_run(child.string.strip())
                apply_styles(run, styles)

    elif tag.name == 'table':
        if tag.find('tr'):
            table = doc.add_table(rows=0, cols=len(tag.find('tr').find_all(['th', 'td'])))
            table.style = 'Table Grid'
            for row in tag.find_all('tr'):
                row_cells = table.add_row().cells
                for i, cell in enumerate(row.find_all(['th', 'td'])):
                    if i < len(row_cells):
                        p = row_cells[i].paragraphs[0]
                        cell_text = cell.get_text(strip=True)
                        if cell_text:
                            run = p.add_run(cell_text)
                            cell_styles = {}
                            if cell.get('style'):
                                cell_styles = {
                                    prop.strip(): val.strip()
                                    for prop, val in (
                                        style.split(':') for style in cell.get('style').split(';') if ':' in style
                                    )
                                }
                            apply_styles(run, cell_styles)

    elif tag.name == 'a' and parent_element:
        add_hyperlink(parent_element, tag.get_text(strip=True), tag.get('href'))
    elif tag.name == 'br':
        if parent_element:
            parent_element.add_run().add_break()
    else:
        # Handle text content or unknown tags
        text_content = tag.get_text(strip=True) if hasattr(tag, 'get_text') else str(tag).strip()
        if text_content and parent_element:
            run = parent_element.add_run(text_content)
            apply_styles(run, styles)
        elif text_content and not parent_element:
            # If no parent element (top-level), create a paragraph
            p = doc.add_paragraph()
            run = p.add_run(text_content)
            apply_styles(run, styles)


def convert_html_to_docx(html_content, docx_filename):
    doc = Document()
    soup = BeautifulSoup(html_content, 'lxml')

    body = soup.find('body')
    if body:
        for element in body.children:
            if element.name:
                handle_tag(element, doc, None)  # Pass None as parent_element for top-level elements
    else:
        # If no body tag, process all elements
        for element in soup.children:
            if element.name:
                handle_tag(element, doc, None)

    doc.save(docx_filename)

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    if file:
        # Generate unique job ID
        job_id = str(uuid.uuid4())
        
        # Save uploaded file
        filename = f"{job_id}_{file.filename}"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        # Create job entry
        jobs[job_id] = {
            'jobId': job_id,
            'status': 'pending',
            'progress': 0,
            'filename': filename,
            'original_filename': file.filename,
            'created_at': datetime.now(),
            'downloadUrl': None,
            'error': None
        }
        
        # Start conversion in background
        threading.Thread(target=process_conversion, args=(job_id,)).start()
        
        return jsonify({'jobId': job_id})

@app.route('/api/status/<job_id>', methods=['GET'])
def get_job_status(job_id):
    if job_id not in jobs:
        return jsonify({'error': 'Job not found'}), 404
    
    job = jobs[job_id]
    return jsonify({
        'jobId': job_id,
        'status': job['status'],
        'progress': job['progress'],
        'downloadUrl': job['downloadUrl'],
        'error': job['error']
    })

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    filepath = os.path.join(OUTPUT_FOLDER, filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    else:
        return jsonify({'error': 'File not found'}), 404

def process_conversion(job_id):
    """Background task to process HTML to DOCX conversion"""
    try:
        job = jobs[job_id]
        
        # Update status to converting
        jobs[job_id]['status'] = 'converting'
        jobs[job_id]['progress'] = 20
        
        # Read HTML file
        input_filepath = os.path.join(UPLOAD_FOLDER, job['filename'])
        with open(input_filepath, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        jobs[job_id]['progress'] = 50
        
        # Convert HTML to DOCX
        output_filename = f"{job_id}_converted.docx"
        output_filepath = os.path.join(OUTPUT_FOLDER, output_filename)
        
        convert_html_to_docx(html_content, output_filepath)
        
        jobs[job_id]['progress'] = 90
        
        # Update job status
        jobs[job_id]['status'] = 'completed'
        jobs[job_id]['progress'] = 100
        jobs[job_id]['downloadUrl'] = f'/api/download/{output_filename}'
        
        # Clean up input file
        if os.path.exists(input_filepath):
            os.remove(input_filepath)
            
    except Exception as e:
        jobs[job_id]['status'] = 'error'
        jobs[job_id]['error'] = str(e)
        print(f"Conversion error for job {job_id}: {e}")

if __name__ == '__main__':
    app.run(debug=True, port=5000)
