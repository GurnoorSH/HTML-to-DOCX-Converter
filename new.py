from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def convert_html_to_docx(html_content, docx_filename):
    # Create a new DOCX document
    doc = Document()

    # Parse HTML content using BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')

    # Define mapping of HTML tags to corresponding DOCX elements
    tag_mapping = {
        'p': (doc.add_paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT),
        'h1': (doc.add_heading, WD_PARAGRAPH_ALIGNMENT.CENTER),
        'h2': (doc.add_heading, WD_PARAGRAPH_ALIGNMENT.CENTER),
        'h3': (doc.add_heading, WD_PARAGRAPH_ALIGNMENT.CENTER),
        'h4': (doc.add_heading, WD_PARAGRAPH_ALIGNMENT.CENTER),
        'h5': (doc.add_heading, WD_PARAGRAPH_ALIGNMENT.CENTER),
        'h6': (doc.add_heading, WD_PARAGRAPH_ALIGNMENT.CENTER),
        'ul': (doc.add_paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT),
        'ol': (doc.add_paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT),
        'li': (doc.add_paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT),
        'a': (doc.add_paragraph, WD_PARAGRAPH_ALIGNMENT.LEFT),
        # Add more mappings for other HTML elements as needed
    }

    for tag in soup.find_all():
        tag_name = tag.name

        if tag_name in tag_mapping:
            # Create the corresponding DOCX element
            element_creator, alignment = tag_mapping[tag_name]
            element = element_creator(tag.get_text())
            
            # Set alignment for the element
            element.alignment = alignment

            # Handling CSS attributes
            if tag.get('style'):
                styles = tag.get('style').split(';')
                for style in styles:
                    property, value = style.split(':')
                    property = property.strip()
                    value = value.strip()
                    if property == 'font-size':
                        element.runs[0].font.size = Pt(float(value[:-2]))
                    elif property == 'font-weight' and value == 'bold':
                        element.runs[0].bold = True
                    elif property == 'font-style' and value == 'italic':
                        element.runs[0].italic = True
                    elif property == 'text-align':
                        if value == 'center':
                            element.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Save the DOCX document
    doc.save(docx_filename)

# Sample HTML content with CSS attributes
html_content = """
<html>
    <body>
        <h1 style="text-align: center; font-size: 24pt;">Centered Heading</h1>
        <p style="text-align: right; font-weight: bold; font-style: italic;">Right-aligned paragraph.</p>
        <ul>
            <li style="font-size: 18pt; font-weight: bold;">Bold item 1</li>
            <li><a href="https://www.example.com" style="color: blue;">Visit Example.com</a></li>
        </ul>
    </body>
</html>
"""

# Convert the HTML to DOCX
convert_html_to_docx(html_content, "output.docx")
